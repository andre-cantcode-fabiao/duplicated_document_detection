from fastapi import HTTPException
from io import BytesIO
from docx import Document
import logging

def extract_text_from_docx(file: bytes) -> str:
    """
    Extracts and returns the text from an uploaded .docx file.

    :param file: UploadFile object from FastAPI
    :return: Extracted text as a single string
    """
    try:
        # Read file content into memory using BytesIO
        doc = Document(BytesIO(file))

        # Extract all text by iterating over paragraphs
        full_text = [paragraph.text for paragraph in doc.paragraphs]

        # Join all paragraphs into a single string
        return '\n'.join(full_text)

    except Exception as e:
        logging.error(f"An error occurred while extracting text: {e}")
        raise HTTPException(status_code=500, detail="Error processing the .docx file from bytes.")

def extract_text_from_docx_path(file_path: str) -> str:
    """
    Extracts and returns the text from a .docx file given its path.

    :param file_path: Path to the .docx file
    :return: Extracted text as a single string
    """
    try:

        doc = Document(file_path)
        full_text = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(full_text)

    except Exception as e:
        print(f"An error occurred while extracting text: {e}")
        raise HTTPException(status_code=500, detail="Error processing the .docx file from path.")